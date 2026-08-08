"""Getting text out of an uploaded resume.

Narrow on purpose. PDF and DOCX only, text layer only, no OCR. A scanned CV
is a real thing people have, and the honest answer is "this cannot read that"
rather than an empty profile the user discovers three emails later.

That failure mode is the reason for the length floor below: an extractor will
happily return a handful of stray characters from a page of images and report
success, and a profile silently built from nothing is worse than an upload
that failed.

PDF extraction goes through pdfplumber with `layout=True` rather than the
obvious `pypdf.extract_text()`, and the reason is worth writing down because
the obvious choice looks fine until you feed it a real CV. Resumes are
overwhelmingly exported from design tools, which place every glyph
individually; pypdf then puts a space between each one and hands the model
"A T H A R V  T I W A R I" - on one real resume, 98% of the tokens it
produced were single characters. The model spends its attention
re-segmenting words instead of reading, and the parse becomes a coin flip.

The same call also fixes reading order. Two-column CVs come out of pypdf in
content-stream order, which scatters the skills block hundreds of lines away
from the projects it belongs to, so per-project tech comes back empty. Keeping
the layout means the column geometry survives as whitespace and the model can
see there are two columns - which is why `_clean_layout` below is careful not
to collapse runs of spaces the way the DOCX path does.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass

MAX_UPLOAD_BYTES = 10 * 1024 * 1024
MAX_PAGES = 20

# Below this, whatever came back is page furniture rather than a CV. A one-page
# resume is comfortably over a thousand characters; 200 is generous.
MIN_MEANINGFUL_CHARS = 200

# Enough for any real CV, and a bound on what gets posted to the model.
MAX_EXTRACTED_CHARS = 40_000

PDF_MAGIC = b"%PDF-"
ZIP_MAGIC = b"PK\x03\x04"


class ResumeError(Exception):
    """Phrased for the person who uploaded the file, not for a log."""


@dataclass(frozen=True)
class Extracted:
    text: str
    pages: int
    kind: str  # pdf | docx


def _clean(text: str) -> str:
    # PDF extraction leaves ragged whitespace and the odd NUL; the model reads
    # this, so tidy it rather than shipping the artefacts as content.
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t ]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()[:MAX_EXTRACTED_CHARS]


def _clean_layout(text: str) -> str:
    """Tidy layout-preserved text without flattening the columns.

    Deliberately does not collapse runs of spaces, which is the one thing
    `_clean` does that would undo the extraction: on a two-column CV those runs
    are all that tells the model the achievements list is a sidebar and not a
    continuation of the summary. Trailing padding carries no such signal, so it
    goes - on a one-page resume that is a fifth of the characters for nothing.
    """
    text = text.replace("\x00", " ")
    lines = [line.rstrip() for line in text.split("\n")]

    kept: list[str] = []
    for line in lines:
        # One blank line separates blocks; a run of them is just page air.
        if not line and (not kept or not kept[-1]):
            continue
        kept.append(line)

    return "\n".join(kept).strip()[:MAX_EXTRACTED_CHARS]


def _extract_pdf(data: bytes) -> Extracted:
    import pdfplumber
    from pdfminer.pdfdocument import PDFPasswordIncorrect
    from pdfminer.pdfparser import PDFSyntaxError

    try:
        pdf = pdfplumber.open(io.BytesIO(data))
    except PDFPasswordIncorrect as exc:
        # pdfplumber retries with an empty password itself, so getting here
        # means a real one rather than "protected but not really".
        raise ResumeError(
            "That PDF is password protected. Remove the password and upload it "
            "again."
        ) from exc
    except (PDFSyntaxError, ValueError, TypeError) as exc:
        raise ResumeError(
            "That PDF could not be opened - it may be corrupted. Try "
            "re-exporting it."
        ) from exc

    with pdf:
        total = len(pdf.pages)
        chunks = []
        for page in pdf.pages[:MAX_PAGES]:
            try:
                # layout=True is what keeps words whole and columns aligned.
                # The flat call is the fallback rather than pypdf because
                # either way pdfplumber groups glyphs into words first, which
                # is the half of the problem that actually broke parsing.
                chunks.append(page.extract_text(layout=True) or page.extract_text() or "")
            except Exception:  # noqa: BLE001 - one bad page should not lose the rest
                continue

    return Extracted(text=_clean_layout("\n".join(chunks)), pages=total, kind="pdf")


def _extract_docx(data: bytes) -> Extracted:
    import docx
    from docx.opc.exceptions import PackageNotFoundError

    try:
        document = docx.Document(io.BytesIO(data))
    except PackageNotFoundError as exc:
        raise ResumeError(
            "That does not look like a .docx. If it is an older .doc, open it "
            "and save it as .docx first."
        ) from exc

    lines = [paragraph.text for paragraph in document.paragraphs]
    # Plenty of CVs lay everything out in a table, and skipping tables would
    # quietly return a name and nothing else.
    for table in document.tables:
        for row in table.rows:
            lines.append(" ".join(cell.text for cell in row.cells))

    return Extracted(text=_clean("\n".join(lines)), pages=0, kind="docx")


def extract_text(data: bytes, filename: str) -> Extracted:
    """Pull the text layer out of an uploaded resume.

    Raises ResumeError with something the user can act on. Never returns an
    empty or near-empty result: that is the case the caller most needs told
    about, and it is indistinguishable from success once it reaches the model.
    """
    if not data:
        raise ResumeError("That file is empty.")
    if len(data) > MAX_UPLOAD_BYTES:
        size = len(data) / (1024 * 1024)
        raise ResumeError(f"That file is {size:.1f} MB; the limit is 10 MB.")

    lowered = filename.lower()
    if lowered.endswith(".pdf") or data.startswith(PDF_MAGIC):
        extracted = _extract_pdf(data)
    elif lowered.endswith(".docx") or data.startswith(ZIP_MAGIC):
        extracted = _extract_docx(data)
    elif lowered.endswith(".doc"):
        raise ResumeError(
            "The old .doc format is not supported. Open it in Word and save it "
            "as .docx, or export a PDF."
        )
    else:
        raise ResumeError(f"Cannot read {filename!r}. Upload a PDF or a .docx.")

    if len(extracted.text) < MIN_MEANINGFUL_CHARS:
        raise ResumeError(
            "There is no selectable text in that file - it is probably a scan "
            "or an image export. This cannot read images, so either export a "
            "text-based PDF from the original document, or fill your profile "
            "in by hand instead. Nothing was saved."
        )

    return extracted
