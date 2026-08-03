"""Resume extraction, parsing, storage and the completeness gate.

The case that matters most here is the one that fails quietly everywhere else:
a scanned CV with no text layer. pypdf returns a few stray characters and
reports success, the model dutifully invents a profile from nothing, and the
user finds out three emails later. So the failure is asserted, not the
happy path alone.
"""

from __future__ import annotations

import asyncio
import io
import sys
import unittest
import uuid
from pathlib import Path
from types import SimpleNamespace

import httpx

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
sys.path.insert(0, str(Path(__file__).resolve().parents[3] / "packages" / "core"))

from app.services.completeness import REQUIRED_SCORE, assess  # noqa: E402
from app.services.gemini import AIError, GeminiClient  # noqa: E402
from app.services.resume import (  # noqa: E402
    MAX_UPLOAD_BYTES, ResumeError, extract_text,
)
from app.services.storage import LocalStorage, StorageError  # noqa: E402

RESUME_LINES = [
    "Dana Sharma",
    "Backend engineer, distributed systems",
    "dana@example.com | github.com/dana | linkedin.com/in/dana",
    "EXPERIENCE",
    "ExampleCorp - Senior Engineer, 2023 to present",
    "Cut p99 latency on the ingest path from 800ms to 90ms.",
    "Owned the migration from a single Postgres to a sharded cluster.",
    "PROJECTS",
    "ratelimit - a token bucket library used in production by four teams.",
    "EDUCATION",
    "BSc Computer Science, University of Somewhere, 2022",
]


def make_pdf(lines: list[str]) -> bytes:
    """A minimal but structurally valid PDF with a real text layer."""
    def escape(text: str) -> str:
        return text.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")

    content = "BT /F1 12 Tf 72 720 Td 14 TL\n"
    content += "".join(f"({escape(line)}) Tj T*\n" for line in lines)
    content += "ET"
    stream = content.encode("latin-1")

    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]

    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for index, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{index} 0 obj\n".encode() + body + b"\nendobj\n"

    xref_at = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for offset in offsets:
        out += f"{offset:010d} 00000 n \n".encode()
    out += f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_at}\n%%EOF\n".encode()
    return bytes(out)


def make_docx(lines: list[str]) -> bytes:
    import docx

    document = docx.Document()
    for line in lines:
        document.add_paragraph(line)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class TestExtraction(unittest.TestCase):
    def test_pdf_text_layer_is_read(self):
        extracted = extract_text(make_pdf(RESUME_LINES), "dana.pdf")
        self.assertEqual(extracted.kind, "pdf")
        self.assertIn("Dana Sharma", extracted.text)
        self.assertIn("ExampleCorp", extracted.text)
        self.assertIn("github.com/dana", extracted.text)

    def test_docx_is_read(self):
        extracted = extract_text(make_docx(RESUME_LINES), "dana.docx")
        self.assertEqual(extracted.kind, "docx")
        self.assertIn("Dana Sharma", extracted.text)
        self.assertIn("ratelimit", extracted.text)

    def test_docx_tables_are_not_skipped(self):
        import docx

        document = docx.Document()
        document.add_paragraph("Dana Sharma")
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "ExampleCorp"
        table.rows[0].cells[1].text = "Senior Engineer, 2023 to present, ingest latency work"
        document.add_paragraph("x" * 200)
        buffer = io.BytesIO()
        document.save(buffer)

        # Plenty of CVs are laid out entirely in a table; skipping them would
        # return a name and nothing else.
        extracted = extract_text(buffer.getvalue(), "table.docx")
        self.assertIn("ExampleCorp", extracted.text)
        self.assertIn("Senior Engineer", extracted.text)

    def test_image_only_pdf_fails_loudly(self):
        # A page with no text operators - what a scan looks like to pypdf.
        with self.assertRaises(ResumeError) as ctx:
            extract_text(make_pdf([]), "scan.pdf")
        message = str(ctx.exception)
        self.assertIn("no selectable text", message)
        self.assertIn("by hand", message)   # points at the manual path
        self.assertIn("Nothing was saved", message)

    def test_nearly_empty_pdf_is_not_treated_as_success(self):
        with self.assertRaises(ResumeError):
            extract_text(make_pdf(["Dana"]), "thin.pdf")

    def test_empty_file_is_rejected(self):
        with self.assertRaises(ResumeError):
            extract_text(b"", "empty.pdf")

    def test_oversized_file_is_rejected_before_parsing(self):
        with self.assertRaises(ResumeError) as ctx:
            extract_text(b"%PDF-" + b"x" * MAX_UPLOAD_BYTES, "huge.pdf")
        self.assertIn("10 MB", str(ctx.exception))

    def test_old_doc_format_says_what_to_do(self):
        with self.assertRaises(ResumeError) as ctx:
            extract_text(b"\xd0\xcf\x11\xe0" + b"x" * 500, "old.doc")
        self.assertIn("save it as .docx", str(ctx.exception))

    def test_unknown_extension_is_rejected(self):
        with self.assertRaises(ResumeError):
            extract_text(b"just some text " * 50, "resume.txt")

    def test_corrupt_pdf_is_rejected(self):
        with self.assertRaises(ResumeError):
            extract_text(b"%PDF-1.4\nnot really a pdf at all\n" * 20, "broken.pdf")

    def test_type_is_sniffed_when_the_name_lies(self):
        # A PDF uploaded as "resume" with no extension still reads.
        extracted = extract_text(make_pdf(RESUME_LINES), "resume")
        self.assertEqual(extracted.kind, "pdf")


# ------------------------------------------------------------------- Gemini


def gemini_returning(payload: dict, status_code: int = 200) -> GeminiClient:
    def handler(_request: httpx.Request) -> httpx.Response:
        return httpx.Response(status_code, json=payload)

    return GeminiClient(api_key="test-key", transport=httpx.MockTransport(handler))


def as_candidate(text: str) -> dict:
    return {"candidates": [{"content": {"parts": [{"text": text}]}}]}


class TestGemini(unittest.TestCase):
    def test_parses_structured_output(self):
        client = gemini_returning(
            as_candidate(
                '{"name":"Dana Sharma","headline":"Backend engineer",'
                '"bio":"Works on distributed systems.","education":"BSc",'
                '"links":{"github":"github.com/dana"},'
                '"projects":[{"name":"ratelimit","summary":"token bucket"}],'
                '"experience":[{"company":"ExampleCorp","role":"Senior Engineer"}]}'
            )
        )
        parsed = asyncio.run(client.parse_resume("Dana Sharma\nBackend engineer"))
        self.assertEqual(parsed["name"], "Dana Sharma")
        self.assertEqual(parsed["projects"][0]["name"], "ratelimit")
        self.assertEqual(parsed["experience"][0]["company"], "ExampleCorp")

    def test_request_carries_the_schema_and_json_mime_type(self):
        captured: dict = {}

        def handler(request: httpx.Request) -> httpx.Response:
            captured.update(__import__("json").loads(request.content))
            captured["headers"] = dict(request.headers)
            captured["url"] = str(request.url)
            return httpx.Response(200, json=as_candidate('{"name":"x"}'))

        client = GeminiClient(api_key="test-key", transport=httpx.MockTransport(handler))
        asyncio.run(client.parse_resume("some resume text"))

        config = captured["generationConfig"]
        self.assertEqual(config["responseMimeType"], "application/json")
        self.assertIn("projects", config["responseSchema"]["properties"])
        # Asking for JSON in the prose is not enough; the schema is what makes
        # the decoder emit conforming output.
        self.assertLessEqual(config["temperature"], 0.2)

    def test_api_key_travels_in_a_header_not_the_url(self):
        captured: dict = {}

        def handler(request: httpx.Request) -> httpx.Response:
            captured["url"] = str(request.url)
            captured["header"] = request.headers.get("x-goog-api-key")
            return httpx.Response(200, json=as_candidate('{"name":"x"}'))

        client = GeminiClient(api_key="super-secret", transport=httpx.MockTransport(handler))
        asyncio.run(client.parse_resume("text"))

        self.assertEqual(captured["header"], "super-secret")
        self.assertNotIn("super-secret", captured["url"])

    def test_extraction_prompt_forbids_invention(self):
        captured: dict = {}

        def handler(request: httpx.Request) -> httpx.Response:
            body = __import__("json").loads(request.content)
            captured["prompt"] = body["contents"][0]["parts"][0]["text"]
            return httpx.Response(200, json=as_candidate('{"name":"x"}'))

        client = GeminiClient(api_key="k", transport=httpx.MockTransport(handler))
        asyncio.run(client.parse_resume("Dana Sharma"))

        prompt = captured["prompt"]
        self.assertIn("Never infer, embellish or fill a gap", prompt)
        self.assertIn("Dana Sharma", prompt)

    def test_missing_key_is_refused_before_any_request(self):
        with self.assertRaises(AIError) as ctx:
            asyncio.run(GeminiClient(api_key="").parse_resume("text"))
        self.assertIn("GEMINI_API_KEY", str(ctx.exception))

    def test_empty_text_is_refused(self):
        with self.assertRaises(AIError):
            asyncio.run(gemini_returning(as_candidate("{}")).parse_resume("   "))

    def test_rejected_key_is_reported_as_such(self):
        client = gemini_returning({"error": {"message": "API key not valid"}}, status_code=403)
        with self.assertRaises(AIError) as ctx:
            asyncio.run(client.parse_resume("text"))
        self.assertIn("rejected the API key", str(ctx.exception))

    def test_rate_limit_is_reported_as_such(self):
        client = gemini_returning({"error": {"message": "quota"}}, status_code=429)
        with self.assertRaises(AIError) as ctx:
            asyncio.run(client.parse_resume("text"))
        self.assertIn("rate limit", str(ctx.exception))

    def test_blocked_prompt_is_reported(self):
        client = gemini_returning({"promptFeedback": {"blockReason": "SAFETY"}})
        with self.assertRaises(AIError) as ctx:
            asyncio.run(client.parse_resume("text"))
        self.assertIn("SAFETY", str(ctx.exception))

    def test_truncated_output_says_so(self):
        client = gemini_returning({"candidates": [{"finishReason": "MAX_TOKENS"}]})
        with self.assertRaises(AIError) as ctx:
            asyncio.run(client.parse_resume("text"))
        self.assertIn("output limit", str(ctx.exception))

    def test_malformed_json_is_not_passed_off_as_a_profile(self):
        client = gemini_returning(as_candidate("here you go: {name: Dana,"))
        with self.assertRaises(AIError) as ctx:
            asyncio.run(client.parse_resume("text"))
        self.assertIn("malformed JSON", str(ctx.exception))

    def test_json_array_is_rejected(self):
        client = gemini_returning(as_candidate("[1, 2, 3]"))
        with self.assertRaises(AIError):
            asyncio.run(client.parse_resume("text"))


# ------------------------------------------------------------------ storage


class TestStorage(unittest.TestCase):
    def setUp(self):
        import tempfile

        self.tmp = tempfile.TemporaryDirectory()
        self.storage = LocalStorage(self.tmp.name)
        self.user = uuid.uuid4()

    def tearDown(self):
        self.tmp.cleanup()

    def test_round_trip(self):
        key = self.storage.key_for(self.user, "dana.pdf")
        self.storage.put(key, b"bytes")
        self.assertEqual(self.storage.get(key), b"bytes")

    def test_delete_really_unlinks(self):
        key = self.storage.key_for(self.user, "dana.pdf")
        self.storage.put(key, b"bytes")
        self.assertTrue(self.storage.delete(key))
        self.assertFalse((Path(self.tmp.name) / key).exists())
        with self.assertRaises(StorageError):
            self.storage.get(key)

    def test_deleting_twice_is_not_an_error(self):
        key = self.storage.key_for(self.user, "dana.pdf")
        self.storage.put(key, b"x")
        self.assertTrue(self.storage.delete(key))
        self.assertFalse(self.storage.delete(key))

    def test_keys_are_scoped_to_the_user(self):
        key = self.storage.key_for(self.user, "dana.pdf")
        self.assertTrue(key.startswith(f"{self.user}/"))

    def test_uploaded_filename_is_not_used_as_the_key(self):
        key = self.storage.key_for(self.user, "../../etc/passwd")
        self.assertNotIn("..", key)
        self.assertNotIn("passwd", key)

    def test_key_escaping_the_root_is_refused(self):
        with self.assertRaises(StorageError):
            self.storage.put("../escaped.pdf", b"x")

    def test_delete_prefix_removes_everything_for_one_user(self):
        for name in ("a.pdf", "b.pdf"):
            self.storage.put(self.storage.key_for(self.user, name), b"x")
        other = uuid.uuid4()
        self.storage.put(self.storage.key_for(other, "c.pdf"), b"x")

        self.assertEqual(self.storage.delete_prefix(str(self.user)), 2)
        self.assertFalse((Path(self.tmp.name) / str(self.user)).exists())
        self.assertTrue((Path(self.tmp.name) / str(other)).exists())


# ------------------------------------------------------------- completeness


def profile(**overrides):
    base = dict(headline="", bio="", education="", availability="", links={})
    base.update(overrides)
    return SimpleNamespace(**base)


class TestCompleteness(unittest.TestCase):
    def test_empty_profile_cannot_create_targets(self):
        result = assess(profile(), [], [])
        self.assertEqual(result.score, 0)
        self.assertFalse(result.complete)
        self.assertTrue(result.blocks_targets)

    def test_missing_fields_come_back_with_something_to_do(self):
        result = assess(profile(), [], [])
        self.assertIn("headline", result.missing)
        self.assertTrue(all(prompt.strip() for prompt in result.prompts))
        self.assertEqual(len(result.missing), len(result.prompts))

    def test_a_usable_profile_passes(self):
        result = assess(
            profile(
                headline="Backend engineer",
                bio="Works on distributed systems.",
                links={"github": "github.com/dana"},
                education="BSc Computer Science",
            ),
            [SimpleNamespace(name="ratelimit")],
            [],
        )
        self.assertGreaterEqual(result.score, REQUIRED_SCORE)
        self.assertTrue(result.complete)

    def test_prose_without_evidence_is_not_enough(self):
        # Headline and bio alone is exactly the profile that produces filler.
        result = assess(
            profile(headline="Backend engineer", bio="Distributed systems.", links={"github": "g"}),
            [], [],
        )
        self.assertFalse(result.complete)
        self.assertIn("evidence", result.missing)

    def test_experience_counts_as_evidence_too(self):
        result = assess(
            profile(
                headline="Backend engineer",
                bio="Distributed systems.",
                links={"github": "g"},
                education="BSc",
            ),
            [],
            [SimpleNamespace(company="ExampleCorp")],
        )
        self.assertTrue(result.complete)

    def test_availability_alone_never_blocks(self):
        full = assess(
            profile(
                headline="h", bio="b", links={"github": "g"},
                education="e", availability="From June",
            ),
            [SimpleNamespace(name="p")], [],
        )
        without = assess(
            profile(headline="h", bio="b", links={"github": "g"}, education="e"),
            [SimpleNamespace(name="p")], [],
        )
        self.assertTrue(full.complete)
        self.assertTrue(without.complete)

    def test_blank_links_do_not_count(self):
        result = assess(profile(links={"github": "   ", "linkedin": ""}), [], [])
        self.assertIn("links", result.missing)


if __name__ == "__main__":
    unittest.main(verbosity=2)
